"""Экспорт сгенерированного договора в DOCX и PDF.
Принимаем текст с лёгким Markdown (заголовки `# H1`, `## H2`, списки, **bold**, *italic*).

DOCX — python-docx, PDF — weasyprint (markdown → HTML → PDF)."""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Tuple


# ───────────────────── helpers ─────────────────────

def _filename_base(contract_type: str | None) -> str:
    safe_type = re.sub(r'[^\w\-]+', '_', (contract_type or 'contract')).strip('_')
    return f"{safe_type or 'contract'}_{datetime.now().strftime('%Y%m%d_%H%M')}"


def _strip_inline_md(text: str) -> str:
    """Убираем `**bold**` / `*italic*` для plain DOCX run (где inline-форматы
    обрабатываются параграфом-уровнем, а не внутри run)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)\*(?!\*)', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


# ───────────────────── DOCX ─────────────────────

def to_docx(text: str, contract_type: str | None = None) -> Tuple[bytes, str]:
    """Возвращает (bytes, filename). Заголовки `# ` / `## ` / `### `,
    маркированные списки `- ` и `* `, нумерованные `1. ` сохраняются."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Базовый стиль
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for sec in doc.sections:
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)

    in_table = False  # простой markdown — таблицы не парсим, оставим как текст

    for raw_line in (text or '').splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Заголовки
        m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if m:
            level = min(len(m.group(1)), 4)
            title = _strip_inline_md(m.group(2))
            heading = doc.add_heading(title, level=level)
            if level == 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in heading.runs:
                    run.bold = True
            continue

        # Маркированный список
        if re.match(r'^[\-\*•]\s+', stripped):
            doc.add_paragraph(_strip_inline_md(re.sub(r'^[\-\*•]\s+', '', stripped)),
                              style='List Bullet')
            continue

        # Нумерованный список
        if re.match(r'^\d+[\.\)]\s+', stripped):
            doc.add_paragraph(_strip_inline_md(re.sub(r'^\d+[\.\)]\s+', '', stripped)),
                              style='List Number')
            continue

        # Обычный параграф с инлайн bold/italic
        p = doc.add_paragraph()
        # Простейший токенайзер для **bold** и *italic*
        tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', stripped)
        for tok in tokens:
            if not tok:
                continue
            if tok.startswith('**') and tok.endswith('**'):
                run = p.add_run(tok[2:-2]); run.bold = True
            elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
                run = p.add_run(tok[1:-1]); run.italic = True
            elif tok.startswith('`') and tok.endswith('`'):
                run = p.add_run(tok[1:-1]); run.font.name = 'Courier New'
            else:
                p.add_run(tok)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), f"{_filename_base(contract_type)}.docx"


# ───────────────────── PDF (reportlab — без системных зависимостей) ─────────────────────

def _markdown_to_html(md: str) -> str:
    """Минимальный markdown → HTML без внешней зависимости."""
    out_lines = []
    in_ul = False
    in_ol = False

    def close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            out_lines.append('</ul>'); in_ul = False
        if in_ol:
            out_lines.append('</ol>'); in_ol = False

    def inline(s: str) -> str:
        s = (s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
        s = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', s)
        s = re.sub(r'(?<!\*)\*(?!\*)(.+?)\*(?!\*)', r'<em>\1</em>', s)
        s = re.sub(r'`([^`]+)`', r'<code>\1</code>', s)
        return s

    for raw in (md or '').splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            close_lists()
            continue

        m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if m:
            close_lists()
            level = min(len(m.group(1)), 4)
            out_lines.append(f'<h{level}>{inline(m.group(2))}</h{level}>')
            continue

        if re.match(r'^[\-\*•]\s+', stripped):
            if not in_ul:
                close_lists(); out_lines.append('<ul>'); in_ul = True
            item = re.sub(r'^[\-\*•]\s+', '', stripped)
            out_lines.append(f'<li>{inline(item)}</li>')
            continue

        if re.match(r'^\d+[\.\)]\s+', stripped):
            if not in_ol:
                close_lists(); out_lines.append('<ol>'); in_ol = True
            item = re.sub(r'^\d+[\.\)]\s+', '', stripped)
            out_lines.append(f'<li>{inline(item)}</li>')
            continue

        close_lists()
        out_lines.append(f'<p>{inline(stripped)}</p>')

    close_lists()
    return '\n'.join(out_lines)


def _inline_to_reportlab(s: str) -> str:
    """`**bold**`/`*italic*`/`` `code` `` → reportlab-mini-html."""
    s = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    s = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', s)
    s = re.sub(r'(?<!\*)\*(?!\*)(.+?)\*(?!\*)', r'<i>\1</i>', s)
    s = re.sub(r'`([^`]+)`', r'<font face="Courier">\1</font>', s)
    return s


def to_pdf(text: str, contract_type: str | None = None) -> Tuple[bytes, str]:
    """Markdown → PDF через reportlab Platypus (поддержка кириллицы из коробки)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    # Подключаем DejaVu — стандартная Unicode-семья, есть в python-docx и в системе
    base_paths = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf',
        '/usr/share/fonts/dejavu/DejaVuSerif.ttf',
    ]
    base_path = next((p for p in base_paths if os.path.exists(p)), None)
    bold_path_candidates = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf',
        '/usr/share/fonts/dejavu/DejaVuSerif-Bold.ttf',
    ]
    italic_path_candidates = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf',
        '/usr/share/fonts/dejavu/DejaVuSerif-Italic.ttf',
    ]
    italic_sans_candidates = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf',
        '/usr/share/fonts/dejavu/DejaVuSans-Oblique.ttf',
    ]
    if base_path:
        try:
            registered = {'normal': 'DejaVuSerif'}
            pdfmetrics.registerFont(TTFont('DejaVuSerif', base_path))
            for p in bold_path_candidates:
                if os.path.exists(p):
                    pdfmetrics.registerFont(TTFont('DejaVuSerif-Bold', p))
                    registered['bold'] = 'DejaVuSerif-Bold'; break
            for p in italic_path_candidates:
                if os.path.exists(p):
                    pdfmetrics.registerFont(TTFont('DejaVuSerif-Italic', p))
                    registered['italic'] = 'DejaVuSerif-Italic'; break
            # DejaVu Serif не имеет italic — берём Sans-Oblique как замену
            if 'italic' not in registered:
                for p in italic_sans_candidates:
                    if os.path.exists(p):
                        pdfmetrics.registerFont(TTFont('DejaVuSerif-Italic', p))
                        registered['italic'] = 'DejaVuSerif-Italic'; break
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            # Передаём только реально зарегистрированные варианты
            registerFontFamily('DejaVuSerif', **registered)
            FONT = 'DejaVuSerif'
        except Exception as e:
            print(f"⚠️ registerFontFamily failed: {e}")
            FONT = 'Helvetica'
    else:
        FONT = 'Helvetica'  # fallback (latin-only — но не упадёт)

    styles = getSampleStyleSheet()
    style_body = ParagraphStyle('body', parent=styles['BodyText'],
                                fontName=FONT, fontSize=11, leading=15,
                                alignment=TA_JUSTIFY, spaceAfter=4)
    style_h1 = ParagraphStyle('h1', parent=styles['Heading1'],
                              fontName=FONT, fontSize=16, leading=20,
                              alignment=TA_CENTER, spaceBefore=4, spaceAfter=12)
    style_h2 = ParagraphStyle('h2', parent=styles['Heading2'],
                              fontName=FONT, fontSize=13, leading=16,
                              spaceBefore=10, spaceAfter=6)
    style_h3 = ParagraphStyle('h3', parent=styles['Heading3'],
                              fontName=FONT, fontSize=12, leading=15,
                              spaceBefore=8, spaceAfter=4)
    style_li = ParagraphStyle('li', parent=style_body, leftIndent=8, spaceAfter=2)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm,
                            title=f"Договор — {contract_type or ''}")

    story = []
    bullet_items: list = []
    number_items: list = []

    def flush_lists():
        nonlocal bullet_items, number_items
        if bullet_items:
            story.append(ListFlowable(
                [ListItem(Paragraph(x, style_li), leftIndent=14) for x in bullet_items],
                bulletType='bullet', start='•'))
            bullet_items = []
        if number_items:
            story.append(ListFlowable(
                [ListItem(Paragraph(x, style_li), leftIndent=14) for x in number_items],
                bulletType='1', start='1'))
            number_items = []

    for raw in (text or '').splitlines():
        line = raw.strip()
        if not line:
            flush_lists()
            story.append(Spacer(1, 4))
            continue

        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            flush_lists()
            level = len(m.group(1))
            content = _inline_to_reportlab(m.group(2))
            if level == 1:
                story.append(Paragraph(content, style_h1))
            elif level == 2:
                story.append(Paragraph(content, style_h2))
            else:
                story.append(Paragraph(content, style_h3))
            continue

        if re.match(r'^[\-\*•]\s+', line):
            flush_lists() if number_items else None
            bullet_items.append(_inline_to_reportlab(re.sub(r'^[\-\*•]\s+', '', line)))
            continue

        if re.match(r'^\d+[\.\)]\s+', line):
            flush_lists() if bullet_items else None
            number_items.append(_inline_to_reportlab(re.sub(r'^\d+[\.\)]\s+', '', line)))
            continue

        flush_lists()
        story.append(Paragraph(_inline_to_reportlab(line), style_body))

    flush_lists()

    doc.build(story)
    buf.seek(0)
    return buf.read(), f"{_filename_base(contract_type)}.pdf"
